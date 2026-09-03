import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

import gradio as gr
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


# ============================================================
# CV STORAGE
# ============================================================

website_cv_text = ""
cv_data = {}


# ============================================================
# CV UPLOAD
# ============================================================

def upload_cv(file):
    global website_cv_text

    if file is None:
        return "⚠️ Please upload your CV."

    filename = file.name

    try:
        if filename.lower().endswith(".pdf"):
            reader = PdfReader(filename)
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

        elif filename.lower().endswith(".docx"):
            document = Document(filename)
            text = ""

            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"

        else:
            return "⚠️ Please upload a PDF or DOCX file."

        if not text.strip():
            return "⚠️ No readable text was found."

        website_cv_text = text

        return (
            "## ✅ CV Uploaded Successfully\n\n"
            f"**File:** {filename.split('/')[-1]}\n\n"
            f"**Characters extracted:** {len(text)}"
        )

    except Exception as error:
        return (
            "⚠️ **Could not read the CV**\n\n"
            f"`{error}`"
        )



def analyse_cv(job_description):

    if not website_cv_text.strip():

        return (
            "⚠️ Please upload your CV first.",
            "",
            "",
            "",
            ""
        )

    if not job_description.strip():

        return (
            "⚠️ Please enter a job description.",
            "",
            "",
            "",
            ""
        )

    cv_lower = website_cv_text.lower()
    job_lower = job_description.lower()

    words = set(
        word.strip(
            ".,!?;:()[]{}"
        )
        for word in job_lower.split()
    )

    words = {
        word for word in words
        if len(word) > 2
    }

    matched = [
        word for word in words
        if word in cv_lower
    ]

    if words:

        match_score = int(
            len(matched) / len(words) * 100
        )

    else:

        match_score = 0

    word_count = len(
        website_cv_text.split()
    )

    if word_count >= 300:

        structure_score = 90

    elif word_count >= 150:

        structure_score = 75

    elif word_count >= 75:

        structure_score = 60

    else:

        structure_score = 40

    skills = [
        "python",
        "sql",
        "excel",
        "power bi",
        "tableau",
        "statistics",
        "data analysis",
        "machine learning",
        "communication",
        "leadership",
        "project management"
    ]

    detected_skills = [
        skill for skill in skills
        if skill in cv_lower
    ]

    skill_score = min(
        100,
        len(detected_skills) * 10
    )

    overall = int(
        structure_score * 0.30
        + match_score * 0.50
        + skill_score * 0.20
    )

    missing = [
        word for word in words
        if word not in cv_lower
    ]

    matching_text = (
        "\n".join(
            f"✓ {word}"
            for word in sorted(matched)[:30]
        )
        if matched
        else "None detected."
    )

    missing_text = (
        "\n".join(
            f"⚠ {word}"
            for word in sorted(missing)[:30]
        )
        if missing
        else "None detected."
    )

    report = f"""
# 📊 CVFix-SA Analysis

## Overall Score: {overall}/100

| Category | Score |
|---|---:|
| 📄 CV Structure | {structure_score}/100 |
| 💼 Job Match | {match_score}/100 |
| 🛠 Skills | {skill_score}/100 |

---

## ✅ Matching Keywords

{matching_text}

---

## ⚠️ Potential Missing Keywords

{missing_text}

---

## 💡 Recommendation

Tailor your CV to the requirements of the job,
but only include skills and experience that are
genuinely true about you.
"""

    return (
        report,
        f"{overall}/100",
        f"{structure_score}/100",
        f"{match_score}/100",
        f"{skill_score}/100"
    )


# ============================================================
# CV IMPROVEMENT
# ============================================================


# ============================================================
# CVFix-SA — SMART FREE CV MATCHING ENGINE
# ============================================================

def smart_normalize(text):
    """
    Normalise text for more reliable CV/job matching.
    """
    text = str(text).lower()

    replacements = {
        "machine-learning": "machine learning",
        "data-analysis": "data analysis",
        "data-analytics": "data analytics",
        "data-engineering": "data engineering",
        "data-science": "data science",
        "powerbi": "power bi",
        "scikit learn": "scikit-learn",
        "scikit-learn": "scikit-learn",
        "r-studio": "r studio",
        "problem-solving": "problem solving",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"[^a-z0-9+#.\s-]", " ", text)
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def smart_keyword_groups():
    """
    Related terminology used by CVFix-SA.
    A job requirement can match a CV through
    closely related professional terminology.
    """

    return {
        "python": [
            "python",
            "python programming",
            "python development",
            "python developer"
        ],

        "sql": [
            "sql",
            "sql queries",
            "sql programming",
            "database queries"
        ],

        "excel": [
            "excel",
            "microsoft excel",
            "spreadsheet",
            "spreadsheets"
        ],

        "power bi": [
            "power bi",
            "powerbi",
            "business intelligence"
        ],

        "tableau": [
            "tableau"
        ],

        "r": [
            "r programming",
            "r studio",
            "rstudio",
            "r language"
        ],

        "statistics": [
            "statistics",
            "statistical",
            "statistical analysis",
            "statistical modelling",
            "statistical modeling"
        ],

        "data analysis": [
            "data analysis",
            "data analyst",
            "data analytics",
            "analytics"
        ],

        "data science": [
            "data science",
            "data scientist"
        ],

        "machine learning": [
            "machine learning",
            "ml",
            "predictive modelling",
            "predictive modeling",
            "predictive analytics"
        ],

        "data engineering": [
            "data engineering",
            "data engineer",
            "etl",
            "data pipeline",
            "data pipelines",
            "data infrastructure"
        ],

        "pandas": [
            "pandas",
            "python pandas"
        ],

        "numpy": [
            "numpy",
            "python numpy"
        ],

        "scikit-learn": [
            "scikit-learn",
            "sklearn",
            "machine learning library"
        ],

        "data visualisation": [
            "data visualisation",
            "data visualization",
            "visualisation",
            "visualization",
            "data presentation",
            "dashboard",
            "charts",
            "graphs"
        ],

        "communication": [
            "communication",
            "communicate",
            "written communication",
            "verbal communication",
            "presentation skills"
        ],

        "problem solving": [
            "problem solving",
            "problem-solving",
            "analytical thinking",
            "critical thinking"
        ],

        "project management": [
            "project management",
            "project manager",
            "project coordination"
        ],

        "research": [
            "research",
            "research project",
            "research methods",
            "research analysis"
        ],

        "reporting": [
            "reporting",
            "reports",
            "report preparation",
            "report writing"
        ],

        "database": [
            "database",
            "databases",
            "database management",
            "data storage"
        ],

        "git": [
            "git",
            "github",
            "version control"
        ],

        "software development": [
            "software development",
            "software developer",
            "application development",
            "programming",
            "coding"
        ],

        "programming": [
            "programming",
            "coding",
            "software development",
            "development"
        ],

        "data cleaning": [
            "data cleaning",
            "data cleansing",
            "data preprocessing",
            "data preparation"
        ],

        "data visualisation": [
            "data visualisation",
            "data visualization",
            "visualisation",
            "visualization",
            "charts",
            "graphs",
            "dashboards"
        ],

        "forecasting": [
            "forecasting",
            "forecast",
            "time series",
            "time-series analysis"
        ],

        "mathematics": [
            "mathematics",
            "mathematical",
            "mathematical sciences"
        ],

        "analytical": [
            "analytical",
            "analysis",
            "analytical skills",
            "analytical thinking"
        ]
    }


def smart_match_keywords(cv_text, job_text):
    """
    Compare CV and job description using related
    terminology rather than exact word matching.
    """

    cv = smart_normalize(cv_text)
    job = smart_normalize(job_text)

    groups = smart_keyword_groups()

    matched = []
    required = []
    missing = []

    for label, variations in groups.items():

        job_found = any(
            variation in job
            for variation in variations
        )

        cv_found = any(
            variation in cv
            for variation in variations
        )

        if job_found:
            required.append(label)

            if cv_found:
                matched.append(label)
            else:
                missing.append(label)

    if required:
        percentage = int(
            len(matched) / len(required) * 100
        )
    else:
        percentage = 0

    return {
        "required": required,
        "matched": matched,
        "missing": missing,
        "percentage": percentage
    }


def smart_recommendations(match_data):
    """
    Produce useful recommendations without inventing
    qualifications or experience.
    """

    recommendations = []

    missing = match_data.get(
        "missing",
        []
    )

    if "python" in missing:
        recommendations.append(
            "Consider highlighting Python projects or "
            "experience if you genuinely have them."
        )

    if "sql" in missing:
        recommendations.append(
            "If you have SQL experience, make SQL queries, "
            "databases or related projects more visible."
        )

    if "data visualisation" in missing:
        recommendations.append(
            "If applicable, highlight dashboards, charts, "
            "Power BI, Tableau or other visualisation work."
        )

    if "machine learning" in missing:
        recommendations.append(
            "If you have machine-learning projects, "
            "make the models and results more visible."
        )

    if "data engineering" in missing:
        recommendations.append(
            "If applicable, highlight ETL, pipelines, "
            "databases or data-engineering projects."
        )

    if "communication" in missing:
        recommendations.append(
            "Consider showing evidence of communication, "
            "presentations or teamwork where truthful."
        )

    if not recommendations:
        recommendations.append(
            "Your CV contains several relevant terms. "
            "Keep the strongest matching skills near the top."
        )

    return recommendations



# ============================================================
# CVFix-SA — OUTPUT QUALITY HELPERS
# ============================================================

def clean_cv_text(text):
    """
    Clean common formatting and punctuation problems
    without changing the factual meaning.
    """

    if not text:
        return ""

    text = str(text)

    # Normalise whitespace
    text = re.sub(r"\s+", " ", text).strip()

    # Fix common professional abbreviations
    replacements = {
        r"\bSql\b": "SQL",
        r"\bsql\b": "SQL",
        r"\bPower bi\b": "Power BI",
        r"\bpower bi\b": "Power BI",
        r"\bTableau\b": "Tableau",
        r"\bPython\b": "Python",
        r"\bNumPy\b": "NumPy",
        r"\bPandas\b": "Pandas",
        r"\bScikit-learn\b": "Scikit-learn",
        r"\bGithub\b": "GitHub",
        r"\bGitHub\b": "GitHub",
        r"\bRstudio\b": "RStudio",
        r"\bR studio\b": "RStudio"
    }

    for pattern, replacement in replacements.items():
        text = re.sub(
            pattern,
            replacement,
            text
        )

    # Remove repeated punctuation
    text = re.sub(r"\.{2,}", ".", text)
    text = re.sub(r",{2,}", ",", text)
    text = re.sub(r";{2,}", ";", text)

    # Remove spaces before punctuation
    text = re.sub(
        r"\s+([,.;:!?])",
        r"\1",
        text
    )

    # Ensure one space after punctuation
    text = re.sub(
        r"([,;:])(?=[A-Za-z])",
        r"\1 ",
        text
    )

    return text.strip()


def clean_cv_data(data):
    """
    Clean the final structured CV data before DOCX generation.
    """

    if not isinstance(data, dict):
        return data

    # Simple text fields
    for field in [
        "name",
        "title",
        "phone",
        "email",
        "location",
        "about"
    ]:

        if field in data:
            data[field] = clean_cv_text(
                data.get(field, "")
            )

    # Education
    for education in data.get(
        "education",
        []
    ):

        if not isinstance(education, dict):
            continue

        for field in [
            "institution",
            "qualification",
            "stream",
            "date",
            "status",
            "description"
        ]:

            if field in education:
                education[field] = clean_cv_text(
                    education.get(field, "")
                )

    # Experience
    for experience in data.get(
        "experience",
        []
    ):

        if not isinstance(experience, dict):
            continue

        for field in [
            "position",
            "organisation",
            "date",
            "description",
            "reference",
            "reference_phone"
        ]:

            if field in experience:
                experience[field] = clean_cv_text(
                    experience.get(field, "")
                )

    # Skills
    if "skills" in data:

        cleaned_skills = []

        for item in data.get("skills", []):

            cleaned = clean_cv_text(str(item))

            if cleaned:
                cleaned_skills.append(cleaned)

        data["skills"] = cleaned_skills

    # Projects
    if "projects" in data:

        cleaned_projects = []

        for project in data.get("projects", []):

            if isinstance(project, dict):

                cleaned_project = dict(project)

                if "name" in cleaned_project:
                    cleaned_project["name"] = clean_cv_text(
                        cleaned_project.get("name", "")
                    )

                if "description" in cleaned_project:
                    cleaned_project["description"] = clean_cv_text(
                        cleaned_project.get("description", "")
                    )

                cleaned_projects.append(cleaned_project)

            else:
                cleaned = clean_cv_text(str(project))

                if cleaned:
                    cleaned_projects.append(cleaned)

        data["projects"] = cleaned_projects

    # Certifications
    if "certifications" in data:

        cleaned_certifications = []

        for item in data.get("certifications", []):

            cleaned = clean_cv_text(str(item))

            if cleaned:
                cleaned_certifications.append(cleaned)

        data["certifications"] = cleaned_certifications

    # References
    if "references" in data:

        cleaned_references = []

        for item in data.get("references", []):

            cleaned = clean_cv_text(str(item))

            if cleaned:
                cleaned_references.append(cleaned)

        data["references"] = cleaned_references

    return data


def improve_cv(job_description):

    global cv_data

    if not website_cv_text.strip():
        return "⚠️ Please upload your CV first."

    if not job_description.strip():
        return "⚠️ Please enter the job description first."

    # --------------------------------------------------------
    # 1. STRUCTURE THE ORIGINAL CV
    # --------------------------------------------------------

    try:
        cv_data = build_complete_cv_data(
            website_cv_text
        )

    except Exception as error:
        return (
            "⚠️ **Could not structure the CV**\n\n"
            f"`{error}`"
        )

    if not cv_data:
        return (
            "⚠️ **Could not extract CV information.**\n\n"
            "Please check that your CV contains readable text."
        )

    # --------------------------------------------------------
    # 2. SMART JOB MATCHING
    # --------------------------------------------------------

    match_data = smart_match_keywords(
        website_cv_text,
        job_description
    )

    required_keywords = match_data.get(
        "required",
        []
    )

    matched_keywords = match_data.get(
        "matched",
        []
    )

    missing_keywords = match_data.get(
        "missing",
        []
    )

    match_percentage = match_data.get(
        "percentage",
        0
    )

    # --------------------------------------------------------
    # 3. PRIORITISE RELEVANT SKILLS
    # --------------------------------------------------------

    original_skills = cv_data.get(
        "skills",
        []
    )

    if not isinstance(original_skills, list):
        original_skills = []

    matched_skills = []
    other_skills = []

    for skill in original_skills:

        skill_text = str(skill).strip()

        if not skill_text:
            continue

        skill_lower = smart_normalize(
            skill_text
        )

        is_relevant = False

        for keyword in matched_keywords:

            variations = smart_keyword_groups().get(
                keyword,
                [keyword]
            )

            if any(
                variation in skill_lower
                or skill_lower in variation
                for variation in variations
            ):
                is_relevant = True
                break

        if is_relevant:
            matched_skills.append(
                skill_text
            )
        else:
            other_skills.append(
                skill_text
            )

    cv_data["skills"] = (
        matched_skills +
        [
            skill
            for skill in other_skills
            if skill not in matched_skills
        ]
    )

    # --------------------------------------------------------
    # 4. STRENGTHEN PROFESSIONAL SUMMARY
    # --------------------------------------------------------

    current_about = str(
        cv_data.get("about", "")
    ).strip()

    title = str(
        cv_data.get("title", "")
    ).strip()

    skills = cv_data.get(
        "skills",
        []
    )

    skill_text = ", ".join(
        str(skill)
        for skill in skills
        if str(skill).strip()
    )

    if not title:

        if "data engineering" in matched_keywords:
            title = "Data Analyst | Data Engineering"

        elif "machine learning" in matched_keywords:
            title = "Data Analyst | Machine Learning"

        elif "data analysis" in matched_keywords:
            title = "Data Analyst"

        elif "software development" in matched_keywords:
            title = "Junior Software Developer"

        else:
            title = "Mathematical Sciences Professional"

    cv_data["title"] = title

    summary_parts = []

    if current_about:
        summary_parts.append(
            current_about.rstrip(". ")
        )

    if matched_keywords:

        readable_keywords = ", ".join(
            keyword.title()
            for keyword in matched_keywords[:8]
        )

        summary_parts.append(
            "Relevant strengths include "
            + readable_keywords
        )

    if skill_text:

        summary_parts.append(
            "Technical skills include "
            + skill_text
        )

    if summary_parts:

        cv_data["about"] = (
            ". ".join(
                part.strip(". ")
                for part in summary_parts
                if part.strip()
            )
            + "."
        )

    # --------------------------------------------------------
    # 5. IMPROVE EXPERIENCE WORDING
    # --------------------------------------------------------

    experience = cv_data.get(
        "experience",
        []
    )

    if isinstance(experience, list):

        for entry in experience:

            description = str(
                entry.get(
                    "description",
                    ""
                )
            ).strip()

            if description:

                improved = description

                replacements = {
                    "responsible for":
                        "managed and delivered",

                    "worked on":
                        "contributed to",

                    "helped with":
                        "supported",

                    "helped":
                        "supported",

                    "used":
                        "utilised",

                    "worked with":
                        "collaborated with",

                    "involved in":
                        "contributed to",

                    "did":
                        "completed"
                }

                for weak, strong in replacements.items():

                    improved = re.sub(
                        rf"\b{re.escape(weak)}\b",
                        strong,
                        improved,
                        flags=re.IGNORECASE
                    )

                entry["description"] = (
                    smart_normalize(improved)
                    .capitalize()
                    + "."
                )

            else:

                position = str(
                    entry.get(
                        "position",
                        ""
                    )
                ).strip()

                organisation = str(
                    entry.get(
                        "organisation",
                        ""
                    )
                ).strip()

                factual_parts = []

                if position:
                    factual_parts.append(
                        f"Worked as {position}"
                    )

                if organisation:
                    factual_parts.append(
                        f"at {organisation}"
                    )

                if matched_keywords:

                    factual_parts.append(
                        "applying relevant analytical "
                        "and technical skills"
                    )

                if factual_parts:

                    entry["description"] = (
                        " ".join(factual_parts)
                        + "."
                    )

    cv_data["experience"] = experience

    # --------------------------------------------------------
    # 6. IMPROVE PROJECT WORDING
    # --------------------------------------------------------

    projects = cv_data.get(
        "projects",
        []
    )

    improved_projects = []

    if isinstance(projects, list):

        for project in projects:

            project_text = str(
                project
            ).strip()

            if not project_text:
                continue

            improved_project = project_text

            replacements = {
                "worked on":
                    "Developed",

                "helped with":
                    "Contributed to",

                "used":
                    "Utilised",

                "involved in":
                    "Contributed to"
            }

            for weak, strong in replacements.items():

                improved_project = re.sub(
                    rf"\b{re.escape(weak)}\b",
                    strong,
                    improved_project,
                    flags=re.IGNORECASE
                )

            improved_projects.append(
                improved_project
            )

    cv_data["projects"] = improved_projects

    # --------------------------------------------------------
    # FINAL CV QUALITY CLEANUP
    # --------------------------------------------------------

    cv_data = clean_cv_data(cv_data)

    # --------------------------------------------------------
    # 7. GENERATE RECOMMENDATIONS
    # --------------------------------------------------------

    recommendations = smart_recommendations(
        match_data
    )

    # --------------------------------------------------------
    # 8. BUILD REPORT
    # --------------------------------------------------------

    result = (
        "# ✨ CVFix-SA Improvement Report\n\n"
    )

    result += (
        "Your CV has been analysed using CVFix-SA's "
        "**free smart CV matching engine**. "
        "Related professional terminology is considered "
        "instead of relying only on exact word matches.\n\n"
    )

    result += "## 📊 Job Match\n\n"

    result += (
        f"### **{match_percentage}%**\n\n"
    )

    result += (
        f"CVFix-SA found **{len(matched_keywords)}** "
        f"relevant skill areas out of "
        f"**{len(required_keywords)}** detected "
        "requirement areas.\n\n"
    )

    # --------------------------------------------------------
    # MATCHED
    # --------------------------------------------------------

    result += "## 🎯 Relevant Keywords Found\n\n"

    if matched_keywords:

        result += "\n".join(
            f"- ✅ {keyword.title()}"
            for keyword in matched_keywords
        )

    else:

        result += (
            "No strong matches were detected."
        )

    result += "\n\n"

    # --------------------------------------------------------
    # MISSING
    # --------------------------------------------------------

    if missing_keywords:

        result += (
            "## ⚠️ Keywords Not Found in Your CV\n\n"
        )

        result += "\n".join(
            f"- {keyword.title()}"
            for keyword in missing_keywords
        )

        result += (
            "\n\nThese are **recommendations only**. "
            "CVFix-SA will not invent qualifications, "
            "skills or experience that you do not have.\n\n"
        )

    # --------------------------------------------------------
    # RECOMMENDATIONS
    # --------------------------------------------------------

    result += "## 💡 CV Recommendations\n\n"

    result += "\n".join(
        f"- {recommendation}"
        for recommendation in recommendations
    )

    result += "\n\n"

    # --------------------------------------------------------
    # EXTRACTED INFORMATION
    # --------------------------------------------------------

    result += (
        "## 📋 Extracted CV Information\n\n"
    )

    result += (
        f"**Name:** {cv_data.get('name', '')}\n\n"
        f"**Professional Title:** {cv_data.get('title', '')}\n\n"
        f"**Location:** {cv_data.get('location', '')}\n\n"
    )

    # --------------------------------------------------------
    # IMPROVEMENTS
    # --------------------------------------------------------

    result += "## ✨ Improvements Made\n\n"

    result += (
        "- Professional summary strengthened\n"
        "- Relevant skills prioritised\n"
        "- Related terminology recognised\n"
        "- Experience wording improved where possible\n"
        "- Project wording improved where possible\n"
        "- Original qualifications and experience preserved\n"
    )

    # --------------------------------------------------------
    # DOWNLOAD
    # --------------------------------------------------------

    result += (
        "\n## 📄 Improved CV\n\n"
        "Your improved CV data is ready.\n\n"
        "Go to the **Download** tab and click "
        "**Generate Improved CV** to create your "
        "professionally formatted DOCX.\n\n"
    )

    result += (
        "---\n\n"
        "✅ **CVFix-SA successfully tailored your CV "
        "using its free smart matching engine.**"
    )

    return result


def build_complete_cv_data(text):
    """
    Parse CV text into the CVFix-SA structured format.
    Section boundaries are detected explicitly so content from
    one section does not spill into another.
    """

    import re

    if not text or not text.strip():
        return {}

    # --------------------------------------------------------
    # CLEAN TEXT
    # --------------------------------------------------------

    lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in text.splitlines()
        if line.strip()
    ]

    # --------------------------------------------------------
    # DEFAULT STRUCTURE
    # --------------------------------------------------------

    data = {
        "name": "",
        "title": "",
        "phone": "",
        "email": "",
        "location": "",
        "about": "",
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "certifications": [],
        "references": []
    }

    # --------------------------------------------------------
    # SECTION HEADINGS
    # --------------------------------------------------------

    section_aliases = {
        "about": {
            "about",
            "profile",
            "professional profile",
            "personal profile",
            "summary",
            "professional summary"
        },

        "education": {
            "education",
            "academic background",
            "qualifications"
        },

        "experience": {
            "experience",
            "work experience",
            "employment history",
            "professional experience"
        },

        "skills": {
            "skills",
            "technical skills",
            "core skills",
            "competencies"
        },

        "projects": {
            "projects",
            "academic projects",
            "personal projects"
        },

        "certifications": {
            "certifications",
            "certificates",
            "licenses"
        },

        "references": {
            "references"
        }
    }

    heading_lookup = {}

    for section, aliases in section_aliases.items():
        for alias in aliases:
            heading_lookup[alias] = section

    # --------------------------------------------------------
    # SPLIT CV INTO SECTIONS
    # --------------------------------------------------------

    sections = {
        "about": [],
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "certifications": [],
        "references": []
    }

    current_section = None

    for line in lines:

        normalized = re.sub(
            r"[^a-z ]",
            "",
            line.lower()
        ).strip()

        if normalized in heading_lookup:
            current_section = heading_lookup[normalized]
            continue

        if current_section:
            sections[current_section].append(line)

    # --------------------------------------------------------
    # CONTACT DETAILS
    # --------------------------------------------------------

    full_text = "\n".join(lines)

    email_match = re.search(
        r"[\w.+-]+@[\w.-]+\.\w+",
        full_text
    )

    if email_match:
        data["email"] = email_match.group(0)

    phone_match = re.search(
        r"(?:\+27|0)\s?\d{2}[\s-]?\d{3}[\s-]?\d{4}",
        full_text
    )

    if phone_match:
        data["phone"] = phone_match.group(0)

    # --------------------------------------------------------
    # LOCATION
    # --------------------------------------------------------

    location_patterns = [
        r"(?:location|address)\s*[:\-]\s*(.+)",
        r"(?:based in)\s+(.+)"
    ]

    for pattern in location_patterns:

        match = re.search(
            pattern,
            full_text,
            re.IGNORECASE
        )

        if match:
            data["location"] = match.group(1).strip()
            break

    # Common South African location fallback
    if not data["location"]:
        location_candidates = [
            "Bellville South, Cape Town",
            "Bellville, Cape Town",
            "Cape Town"
        ]

        for location in location_candidates:
            if location.lower() in full_text.lower():
                data["location"] = location
                break

    # --------------------------------------------------------
    # NAME
    # --------------------------------------------------------

    ignored_lines = {
        "cv",
        "resume",
        "curriculum vitae",
        "curriculum vitae cv"
    }

    for line in lines[:10]:

        lower = line.lower()

        if lower in ignored_lines:
            continue

        if "@" in line:
            continue

        if re.search(r"(?:\+27|^0\d{9})", line.replace(" ", "")):
            continue

        if len(line.split()) >= 2 and len(line) <= 80:
            data["name"] = line
            break

    # --------------------------------------------------------
    # TITLE
    # --------------------------------------------------------

    title_keywords = [
        "developer",
        "analyst",
        "engineer",
        "scientist",
        "programmer",
        "mathematician",
        "statistician",
        "designer",
        "manager",
        "intern",
        "student"
    ]

    for line in lines[:15]:

        if line == data["name"]:
            continue

        lower = line.lower()

        if "@" in line:
            continue

        if any(keyword in lower for keyword in title_keywords):
            data["title"] = line
            break

    # --------------------------------------------------------
    # ABOUT
    # --------------------------------------------------------

    if sections["about"]:
        data["about"] = " ".join(sections["about"])

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    education_text = sections["education"]

    if education_text:

        education_joined = " ".join(education_text)

        institution = ""
        qualification = ""
        stream = ""
        date = ""
        status = ""

        # Institution
        for candidate in [
            "Cape Peninsula University of Technology",
            "CPUT"
        ]:
            if candidate.lower() in education_joined.lower():
                institution = candidate
                break

        # Qualification
        qualification_patterns = [
            r"(Diploma in [A-Za-z ]+)",
            r"(Bachelor(?:'s)? [A-Za-z ]+)",
            r"(Degree in [A-Za-z ]+)"
        ]

        for pattern in qualification_patterns:
            match = re.search(
                pattern,
                education_joined,
                re.IGNORECASE
            )

            if match:
                qualification = match.group(1).strip()

                # Remove institution accidentally captured
                for institution_name in [
                    "Cape Peninsula University of Technology",
                    "CPUT"
                ]:
                    qualification = re.sub(
                        r"\s*" + re.escape(institution_name) + r"\s*$",
                        "",
                        qualification,
                        flags=re.IGNORECASE
                    ).strip()

                break

        # Stream
        stream_match = re.search(
            r"(Stream:\s*[A-Za-z ]+)",
            education_joined,
            re.IGNORECASE
        )

        if stream_match:
            stream = stream_match.group(1).strip()

        # Dates
        date_match = re.search(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}\s*[-–]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4})",
            education_joined,
            re.IGNORECASE
        )

        if date_match:
            date = date_match.group(1).strip()

        if re.search(
            r"\bcompleted\b",
            education_joined,
            re.IGNORECASE
        ):
            status = "Completed"

        description_parts = []

        if stream:
            description_parts.append(stream)

        education_item = {
            "institution": institution,
            "qualification": qualification,
            "stream": stream,
            "date": date,
            "status": status,
            "description": " ".join(description_parts)
        }

        data["education"].append(education_item)

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    experience_lines = sections["experience"]

    if experience_lines:

        experience_text = " ".join(experience_lines)

        position = ""
        organisation = ""
        date = ""
        description = ""

        # First line is commonly the position
        if experience_lines:
            position = experience_lines[0]

        # Organisation
        if len(experience_lines) > 1:
            organisation = experience_lines[1]

        # Date
        date_match = re.search(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}\s*[-–]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s?\d{4}\s*[-–]\s*(?:Present|Current)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s*[-–]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s*\d{4})",
            experience_text,
            re.IGNORECASE
        )

        if date_match:
            date = date_match.group(1).strip()

        if len(experience_lines) > 2:
            description_parts = experience_lines[2:]

            # Remove standalone location/date information
            cleaned_description = []

            for item in description_parts:
                item_clean = item.strip()

                if item_clean.lower() in {
                    "cape town",
                    "bellville",
                    "bellville south",
                    "cape town, south africa"
                }:
                    continue

                if re.fullmatch(
                    r"(?:\d{4}\s*[-–]\s*(?:\d{4}|present|current))",
                    item_clean,
                    re.IGNORECASE
                ):
                    continue

                cleaned_description.append(item_clean)

            description = " ".join(cleaned_description)

        data["experience"].append({
            "position": position,
            "organisation": organisation,
            "date": date,
            "description": description,
            "reference": "",
            "reference_phone": ""
        })

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    skills = []

    for line in sections["skills"]:

        parts = re.split(
            r"[,|•;]",
            line
        )

        for part in parts:

            skill = part.strip()

            if skill and skill not in skills:
                skills.append(skill)

    data["skills"] = skills

    # --------------------------------------------------------
    # PROJECTS
    # --------------------------------------------------------

    projects = []
    project_lines = sections["projects"]

    # Pair project names with their descriptions.
    # A project normally consists of two consecutive lines:
    # project name
    # project description

    i = 0

    while i < len(project_lines):

        project_name = project_lines[i].strip()

        if not project_name:
            i += 1
            continue

        project_description = ""

        if i + 1 < len(project_lines):
            project_description = project_lines[i + 1].strip()

        projects.append({
            "name": project_name,
            "description": project_description
        })

        i += 2

    data["projects"] = projects

    # --------------------------------------------------------
    # CERTIFICATIONS
    # --------------------------------------------------------

    data["certifications"] = sections["certifications"]

    # --------------------------------------------------------
    # REFERENCES
    # --------------------------------------------------------

    data["references"] = sections["references"]

    return data

def set_cell_shading(cell, fill="E6E6E6"):

    tcPr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)

    tcPr.append(shd)


def add_section_heading(document, title):

    table = document.add_table(
        rows=1,
        cols=1
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)

    set_cell_shading(cell)

    paragraph = cell.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(title.upper())

    run.bold = True
    run.font.size = Pt(10)


def add_body_paragraph(document, text):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)

    run.font.size = Pt(9)


def add_education_entry(
    document,
    institution,
    qualification,
    stream,
    date,
    status,
    description
):

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    left.width = Inches(5.1)
    right.width = Inches(1.8)

    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    paragraph = left.paragraphs[0]

    run = paragraph.add_run(
        institution.upper()
    )

    run.bold = True
    run.font.size = Pt(9.5)

    paragraph = left.add_paragraph()

    run = paragraph.add_run(
        qualification
    )

    run.bold = True
    run.font.size = Pt(8.5)

    if stream:

        paragraph = left.add_paragraph()

        run = paragraph.add_run(stream)

        run.font.size = Pt(8)

    paragraph = left.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(description)

    run.font.size = Pt(8.5)

    paragraph = right.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run(date)

    run.bold = True
    run.font.size = Pt(8.5)

    if status:

        paragraph = right.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run(
            f"({status})"
        )

        run.font.size = Pt(8)


def add_experience_entry(
    document,
    position,
    organisation,
    date,
    description,
    reference=None,
    reference_phone=None
):

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    left.width = Inches(5.1)
    right.width = Inches(1.8)

    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    paragraph = left.paragraphs[0]

    run = paragraph.add_run(
        f"{position.upper()} | {organisation.upper()}"
    )

    run.bold = True
    run.font.size = Pt(9.5)

    paragraph = right.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run(date)

    run.bold = True
    run.font.size = Pt(8.5)

    paragraph = left.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(description)

    run.font.size = Pt(8.5)

    if reference:

        paragraph = left.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(5)

        reference_text = f"Reference: {reference}"

        if reference_phone:
            reference_text += f" | Cell: {reference_phone}"

        run = paragraph.add_run(reference_text)

        run.italic = True
        run.font.size = Pt(7.5)


def create_download():

    if not cv_data:

        return None

    filename = "CVFix-SA_Improved_CV.docx"

    document = Document()

    # --------------------------------------------------------
    # PAGE SETUP
    # --------------------------------------------------------

    section = document.sections[0]

    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # --------------------------------------------------------
    # DEFAULT FONT
    # --------------------------------------------------------

    styles = document.styles

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    # --------------------------------------------------------
    # HEADER
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(
        cv_data["name"].upper()
    )

    run.bold = True
    run.font.size = Pt(22)

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)

    run = paragraph.add_run(
        cv_data["title"]
    )

    run.bold = True
    run.font.size = Pt(10)

    # --------------------------------------------------------
    # CONTACT DETAILS
    # --------------------------------------------------------

    contact_table = document.add_table(
        rows=1,
        cols=3
    )

    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    contacts = [
        "☎ " + cv_data["phone"],
        "✉ " + cv_data["email"],
        "📍 " + cv_data["location"]
    ]

    for i, contact in enumerate(contacts):

        cell = contact_table.cell(0, i)

        paragraph = cell.paragraphs[0]

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(contact)

        run.font.size = Pt(8.5)

    # --------------------------------------------------------
    # ABOUT ME
    # --------------------------------------------------------

    add_section_heading(
        document,
        "About Me"
    )

    add_body_paragraph(
        document,
        cv_data["about"]
    )

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    add_section_heading(
        document,
        "Education"
    )

    for education in cv_data["education"]:

        add_education_entry(
            document,
            education["institution"],
            education["qualification"],
            education["stream"],
            education["date"],
            education["status"],
            education["description"]
        )

    # --------------------------------------------------------
    # WORK EXPERIENCE
    # --------------------------------------------------------

    add_section_heading(
        document,
        "Work Experience"
    )

    for experience in cv_data["experience"]:

        add_experience_entry(
            document,
            experience["position"],
            experience["organisation"],
            experience["date"],
            experience["description"],
            experience["reference"],
            experience["reference_phone"]
        )

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    add_section_heading(
        document,
        "Skills"
    )

    skills = cv_data["skills"]

    skill_table = document.add_table(
        rows=3,
        cols=3
    )

    skill_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    index = 0

    for row in range(3):

        for col in range(3):

            cell = skill_table.cell(row, col)

            if index < len(skills):

                paragraph = cell.paragraphs[0]

                run = paragraph.add_run(
                    "• " + skills[index]
                )

                run.font.size = Pt(8.5)

                index += 1

    # --------------------------------------------------------
    # FOOTER
    # --------------------------------------------------------

    footer = section.footer

    paragraph = footer.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "CVFix-SA • Your CV. Your Career."
    )

    run.font.size = Pt(7)
    run.italic = True

    # --------------------------------------------------------
    # SAVE
    # --------------------------------------------------------

    document.save(filename)

    return filename


# ============================================================
# CSS
# ============================================================

css = """
/* CVFix-SA custom colours */
:root {
    --cvfix-dark-green: #14532D;
    --cvfix-dark-brown: #5C4033;
}

/* Dark green accent areas */
.cvfix-green,
.cvfix-green * {
    border-color: #14532D !important;
}

/* Dark brown tab/section accent */
.cvfix-brown,
.cvfix-brown * {
    border-color: #5C4033 !important;
}

/* Gradio tab selected indicator */
button.selected,
button[aria-selected="true"] {
    border-bottom-color: #5C4033 !important;
}

/* Selected tab underline */
.tabitem.selected {
    border-top-color: #5C4033 !important;
}

/* Hide Gradio default branding */
footer {
    display: none !important;
}

.gradio-container footer {
    display: none !important;
}


.gradio-container {
    max-width: 1150px !important;
    margin: auto !important;
}

.cvfix-hero {
    text-align: center;
    padding: 60px 25px;
    margin-bottom: 35px;
    border-radius: 25px;
    background: linear-gradient(
        135deg,
        #12355B,
        #1F6AA5
    );
    color: white;
}

.cvfix-logo {
    font-size: 52px;
    font-weight: 800;
}

.cvfix-tagline {
    font-size: 25px;
    margin-top: 10px;
    font-weight: 600;
}

.cvfix-description {
    max-width: 700px;
    margin: 18px auto 0 auto;
    font-size: 17px;
    line-height: 1.6;
}

"""



# ============================================================
# CVFIX-SA SEO METADATA
# ============================================================

cvfix_seo = """
<meta name="google-site-verification" content="Nyz8jJ9FEv8NCKHYVmXrzWqCksVPm6uZrJKQkgdR_vo" />

<meta name="description"
content="CVFix-SA is a smart CV improvement and job matching
application founded by Junior Software Developer Mr BF Manikela.
Analyse your CV, match it with job requirements, improve your
CV and generate a professional document.">

<meta name="keywords"
content="CVFix-SA, CV improvement, CV optimisation, CV optimizer,
resume improvement, job matching, CV analysis, job applications,
South Africa, Junior Software Developer Mr BF Manikela">

<meta name="author"
content="Junior Software Developer Mr BF Manikela">

<meta property="og:title"
content="CVFix-SA | Smart CV Improvement & Job Matching">

<meta property="og:description"
content="Improve your CV, match your skills to job requirements,
and generate a professional CV with CVFix-SA.">

<meta property="og:type"
content="website">

<meta name="robots"
content="index, follow">

<meta name="application-name"
content="CVFix-SA">

<script type="application/ld+json">
{
    "@context": "https://schema.org",
    "@type": "WebApplication",
    "name": "CVFix-SA",
    "description": "Smart CV improvement and job matching application.",
    "applicationCategory": "BusinessApplication",
    "creator": {
        "@type": "Person",
        "name": "Junior Software Developer Mr BF Manikela"
    }
}
</script>
"""


# ============================================================
# WEBSITE
# ============================================================

with gr.Blocks(
    title="CVFix-SA | Smart CV Improvement & Job Matching"
) as app:

    # ========================================================
    # HERO
    # ========================================================

    gr.HTML(
        """
        <div class="cvfix-hero">

            
            <div class="cvfix-logo">
                📄 CVFix-SA
            </div>
    

            <h1>
                Smart CV Improvement & Job Matching
            </h1>

            <p>
                Build a stronger CV, match your skills to job
                requirements, and create a professional CV ready
                for your next application.
            </p>

            <div style="
                margin-top: 18px;
                font-size: 15px;
                opacity: 0.95;
            ">
                <strong>Founded by Junior Software Developer Mr BF Manikela</strong><br>
                Junior Software Developer
            </div>

        </div>
        """
    )

    # ========================================================
    # QUICK INTRODUCTION
    # ========================================================

    gr.Markdown(
        """
        ## 🚀 Welcome to CVFix-SA

        **CVFix-SA** is a CV improvement and job-matching
        application designed to help job seekers present their
        skills more effectively.

        ### How it works

        **📄 Upload CV**
        → **💼 Add Job Description**
        → **🔍 Analyse**
        → **✨ Improve**
        → **⬇️ Download**

        CVFix-SA analyses the information in your CV and compares
        it with the requirements of the job description.
        """
    )

    gr.Markdown("---")

    # ========================================================
    # MAIN WORKFLOW
    # ========================================================

    with gr.Tabs():

        # ----------------------------------------------------
        # STEP 1 — UPLOAD
        # ----------------------------------------------------

        with gr.Tab("📄 1. Upload CV"):

            gr.Markdown(
                """
                # 📄 Upload Your CV

                Upload your current CV so CVFix-SA can analyse
                its structure, skills, education, experience and
                projects.

                **Supported formats:** PDF and DOCX
                """
            )

            cv_file = gr.File(
                label="Choose your CV",
                file_types=[".pdf", ".docx"],
                type="filepath"
            )

            upload_result = gr.Markdown()

            cv_file.change(
                fn=upload_cv,
                inputs=cv_file,
                outputs=upload_result
            )

        # ----------------------------------------------------
        # STEP 2 — ANALYSE
        # ----------------------------------------------------

        with gr.Tab("🔍 2. Analyse"):

            gr.Markdown(
                """
                # 🔍 Analyse Your CV

                Paste the job description for the position you
                want to apply for.

                CVFix-SA will compare your CV with the job
                requirements and identify matching and missing
                skills.
                """
            )

            job_description = gr.Textbox(
                label="💼 Job Description",
                placeholder=(
                    "Paste the complete job description here..."
                ),
                lines=15
            )

            analyse_button = gr.Button(
                "🔍 Analyse My CV",
                variant="primary"
            )

            analysis_result = gr.Markdown()

            with gr.Row():

                overall = gr.Textbox(
                    label="🏆 Overall",
                    interactive=False
                )

                structure = gr.Textbox(
                    label="📄 Structure",
                    interactive=False
                )

                match = gr.Textbox(
                    label="💼 Job Match",
                    interactive=False
                )

                skills = gr.Textbox(
                    label="🛠 Skills",
                    interactive=False
                )

            analyse_button.click(
                fn=analyse_cv,
                inputs=job_description,
                outputs=[
                    analysis_result,
                    overall,
                    structure,
                    match,
                    skills
                ]
            )

        # ----------------------------------------------------
        # STEP 3 — IMPROVE
        # ----------------------------------------------------

        with gr.Tab("✨ 3. Improve"):

            gr.Markdown(
                """
                # ✨ Improve Your CV

                CVFix-SA will use the analysis to improve the
                wording and presentation of your CV.

                **Important:** CVFix-SA does not invent skills,
                qualifications or experience that are not supported
                by your original CV.
                """
            )

            improve_button = gr.Button(
                "✨ Improve My CV",
                variant="primary"
            )

            improvement_result = gr.Markdown()

            improve_button.click(
                fn=improve_cv,
                inputs=job_description,
                outputs=improvement_result
            )

        # ----------------------------------------------------
        # STEP 4 — DOWNLOAD
        # ----------------------------------------------------

        with gr.Tab("⬇️ 4. Download"):

            gr.Markdown(
                """
                # ⬇️ Download Your Improved CV

                Once your CV has been improved, generate the
                professional CV document using the CVFix-SA
                template.
                """
            )

            download_button = gr.Button(
                "📄 Generate Professional CV",
                variant="primary"
            )

            download_file = gr.File(
                label="Your CVFix-SA Professional CV"
            )

            download_button.click(
                fn=create_download,
                outputs=download_file
            )

            gr.Markdown(
                """
                ### 🎯 Ready for your next application?

                Your improved CV is structured for professional
                presentation and can be downloaded as a DOCX file.
                """
            )

    # ========================================================
    # ABOUT CVFIX-SA
    # ========================================================

    gr.Markdown("---")

    gr.HTML(
        """
        <div style="
            text-align: center;
            padding: 25px 10px;
            margin-top: 10px;
        ">

            <h2>CVFix-SA</h2>

            <p>
                <strong>Your CV. Your Career.</strong>
            </p>

            <p>
                Smart CV improvement and job matching.
            </p>

            <p style="margin-top: 15px;">
                <strong>Founder & Developer</strong><br>
                Junior Software Developer Mr BF Manikela
            </p>

            <p style="
                margin-top: 20px;
                font-size: 13px;
                opacity: 0.75;
            ">
                © 2026 CVFix-SA. All rights reserved.
            </p>

        </div>
        """
    )


app.launch(
    server_name="0.0.0.0",
    server_port=int(os.environ.get("PORT", 7860)),
    share=False,
    css=css,
    head=cvfix_seo
)
